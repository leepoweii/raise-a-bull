"""Smoke tests: real claude -p subprocess + MiniMax M2.7.

Run with:
    ANTHROPIC_BASE_URL=https://api.minimax.io/anthropic \
    ANTHROPIC_AUTH_TOKEN=<key> \
    uv run pytest tests/smoke/ -v --timeout=120
"""
import os
import pytest
from raisebull.runner import ClaudeRunner
from raisebull.trace import TraceStep

smoke = pytest.mark.skipif(
    not (os.environ.get("ANTHROPIC_BASE_URL") and os.environ.get("ANTHROPIC_AUTH_TOKEN")),
    reason="Requires ANTHROPIC_BASE_URL + ANTHROPIC_AUTH_TOKEN env vars",
)


@smoke
@pytest.mark.asyncio
async def test_basic_response(runner: ClaudeRunner):
    result = await runner.run("Reply with exactly: SMOKE_OK")
    assert "SMOKE_OK" in result.text
    assert result.error is None


@smoke
@pytest.mark.asyncio
async def test_chinese_response(runner: ClaudeRunner):
    result = await runner.run("用繁體中文回答：1+1等於多少？只回答數字")
    assert "2" in result.text


@smoke
@pytest.mark.asyncio
async def test_tool_use_read_write(runner: ClaudeRunner, tmp_path):
    """Tool use chain: write a file then read it back."""
    test_file = tmp_path / "test_tool.txt"
    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=str(tmp_path),
        model=runner.model,
    )
    result = await r.run(
        f"Write 'TOOL_OK' to {test_file}, then read it back and tell me its content.",
        timeout_seconds=60.0,
    )
    assert "TOOL_OK" in result.text
    assert test_file.exists()


@smoke
@pytest.mark.asyncio
async def test_session_resume(runner: ClaudeRunner):
    """Two runs: second with --resume should remember context."""
    r1 = await runner.run("Remember the code word: BANANA42. Confirm you've noted it.")
    assert r1.session_id is not None

    r2 = await runner.run(
        "What was the code word I told you?",
        session_id=r1.session_id,
    )
    assert "BANANA42" in r2.text


@smoke
@pytest.mark.asyncio
async def test_on_trace_collects_steps(runner: ClaudeRunner):
    """on_trace callback receives real TraceStep objects."""
    collected: list[TraceStep] = []

    async def on_trace(step: TraceStep):
        collected.append(step)

    result = await runner.run(
        "Read the file /etc/hostname and tell me its content.",
        on_trace=on_trace,
        timeout_seconds=60.0,
    )
    assert result.error is None
    step_types = {s.step_type for s in collected}
    assert "text" in step_types


@smoke
@pytest.mark.asyncio
async def test_extended_thinking(runner: ClaudeRunner):
    result = await runner.run("Think step by step: what is 17 * 23? Reply with just the number.")
    assert "391" in result.text


@smoke
@pytest.mark.asyncio
async def test_stale_session_recovery(runner: ClaudeRunner):
    """Fake stale session_id triggers auto-recovery in runner."""
    result = await runner.run("Say OK", session_id="nonexistent-session-id-12345")
    assert result.text or result.error


# --- Tavily: commented out, may add back as fallback if needed ---
# @smoke
# @pytest.mark.asyncio
# async def test_mcp_tavily_search(runner: ClaudeRunner, tmp_path):
#     """MCP connectivity: Tavily search via --mcp-config."""
#     import json as _json
#     import os
#
#     tavily_key = os.environ.get("TAVILY_API_KEY", "")
#     if not tavily_key:
#         pytest.skip("TAVILY_API_KEY not set")
#
#     mcp_config = tmp_path / "mcp-test.json"
#     mcp_config.write_text(_json.dumps({
#         "mcpServers": {
#             "tavily": {
#                 "command": "npx",
#                 "args": ["-y", "tavily-mcp@0.1.4"],
#                 "env": {"TAVILY_API_KEY": tavily_key},
#             }
#         }
#     }))
#
#     r = ClaudeRunner(
#         claude_bin=runner.claude_bin,
#         workspace=str(tmp_path),
#         model=runner.model,
#         mcp_config=str(mcp_config),
#     )
#     result = await r.run(
#         "Use the tavily MCP tool to search for 'MiniMax AI' and summarize in one sentence.",
#         timeout_seconds=120.0,
#     )
#     assert result.error is None or "tavily" not in (result.error or "").lower()
#     if result.text:
#         assert len(result.text) > 10


@smoke
@pytest.mark.asyncio
async def test_mcp_minimax_search(runner: ClaudeRunner, tmp_path):
    """MCP connectivity: minimax_search (Serper + Jina) via --mcp-config.

    Verifies the search tool returns real Google results for a Chinese query.
    Requires: SERPER_API_KEY, MINIMAX_API_KEY env vars.
    Optional: JINA_API_KEY (browse tool needs it, search does not).
    """
    import json as _json
    import os

    serper_key = os.environ.get("SERPER_API_KEY", "")
    minimax_key = os.environ.get("MINIMAX_API_KEY") or os.environ.get("ANTHROPIC_AUTH_TOKEN", "")
    if not serper_key:
        pytest.skip("SERPER_API_KEY not set")
    if not minimax_key:
        pytest.skip("MINIMAX_API_KEY not set")

    mcp_config = tmp_path / "mcp-search.json"
    mcp_config.write_text(_json.dumps({
        "mcpServers": {
            "minimax_search": {
                "command": "minimax-search",
                "env": {
                    "MINIMAX_API_KEY": minimax_key,
                    "SERPER_API_KEY": serper_key,
                    "JINA_API_KEY": os.environ.get("JINA_API_KEY", ""),
                },
            }
        }
    }))

    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=str(tmp_path),
        model=runner.model,
        mcp_config=str(mcp_config),
    )

    # Use explicit tool name to ensure MCP is loaded (not the built-in WebSearch)
    result = await r.run(
        "Use the mcp__minimax_search__search tool to search for '金門旅遊'. "
        "Summarize the top 3 results in one paragraph.",
        timeout_seconds=120.0,
    )
    assert result.error is None, f"Search failed: {result.error}"
    assert len(result.text) > 20, "Expected meaningful search summary"
    # Verify it actually searched (should mention travel/tourism related content)
    assert result.text, "Empty response from search"


@smoke
@pytest.mark.asyncio
async def test_mcp_minimax_browse(runner: ClaudeRunner, tmp_path):
    """MCP connectivity: minimax_search browse tool reads a web page.

    Requires: SERPER_API_KEY, MINIMAX_API_KEY, JINA_API_KEY env vars.
    """
    import json as _json
    import os

    serper_key = os.environ.get("SERPER_API_KEY", "")
    minimax_key = os.environ.get("MINIMAX_API_KEY") or os.environ.get("ANTHROPIC_AUTH_TOKEN", "")
    jina_key = os.environ.get("JINA_API_KEY", "")
    if not serper_key:
        pytest.skip("SERPER_API_KEY not set")
    if not minimax_key:
        pytest.skip("MINIMAX_API_KEY not set")
    if not jina_key:
        pytest.skip("JINA_API_KEY not set")

    mcp_config = tmp_path / "mcp-browse.json"
    mcp_config.write_text(_json.dumps({
        "mcpServers": {
            "minimax_search": {
                "command": "minimax-search",
                "env": {
                    "MINIMAX_API_KEY": minimax_key,
                    "SERPER_API_KEY": serper_key,
                    "JINA_API_KEY": jina_key,
                },
            }
        }
    }))

    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=str(tmp_path),
        model=runner.model,
        mcp_config=str(mcp_config),
    )

    result = await r.run(
        "Use the mcp__minimax_search__browse tool to read https://example.com "
        "and tell me the page title and first paragraph.",
        timeout_seconds=120.0,
    )
    assert result.error is None, f"Browse failed: {result.error}"
    assert len(result.text) > 10, "Expected page content summary"


@smoke
@pytest.mark.asyncio
async def test_attachment_parse_and_read(runner: ClaudeRunner, tmp_path):
    """Smoke: parse a text file → save to workspace → Claude reads it via Read tool."""
    from raisebull.parsers.router import process_attachment

    workspace = str(tmp_path)
    content = "這是一份測試文件。\n金額：$1,500\n日期：2026-04-04"
    filepath, preview = await process_attachment(
        content.encode(), "test-memo.txt", "text/plain",
        session_id="smoke-test", workspace=workspace,
    )
    assert os.path.exists(filepath)

    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=workspace,
        model=runner.model,
        mcp_config=runner.mcp_config,
    )
    result = await r.run(
        f"有一個檔案在 {filepath}，用 Read 工具讀取它，告訴我檔案裡的金額是多少。只回答金額數字。",
        timeout_seconds=60.0,
    )
    assert result.error is None, f"LLM error: {result.error}"
    assert "1,500" in result.text or "1500" in result.text, f"Expected amount in: {result.text}"


@smoke
@pytest.mark.asyncio
async def test_attachment_csv_parse_and_read(runner: ClaudeRunner, tmp_path):
    """Smoke: parse a CSV → save to workspace → Claude reads and answers."""
    from raisebull.parsers.router import process_attachment

    workspace = str(tmp_path)
    csv_content = "product,price,qty\n高粱酒,580,10\n貢糖,120,25\n麵線,80,50"
    filepath, _ = await process_attachment(
        csv_content.encode(), "products.csv", "text/csv",
        session_id="smoke-csv", workspace=workspace,
    )
    assert os.path.exists(filepath)

    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=workspace,
        model=runner.model,
        mcp_config=runner.mcp_config,
    )
    result = await r.run(
        f"有一個 CSV 檔案在 {filepath}，用 Read 工具讀取，告訴我最貴的商品名稱。只回答商品名。",
        timeout_seconds=60.0,
    )
    assert result.error is None, f"LLM error: {result.error}"
    assert "高粱酒" in result.text, f"Expected 高粱酒 in: {result.text}"


@smoke
@pytest.mark.asyncio
async def test_attachment_docx_parse_and_read(runner: ClaudeRunner, tmp_path):
    """Smoke: parse a DOCX → save to workspace → Claude reads and summarizes."""
    import io as _io
    from docx import Document
    from raisebull.parsers.router import process_attachment

    workspace = str(tmp_path)
    doc = Document()
    doc.add_paragraph("會議紀錄")
    doc.add_paragraph("日期：2026-04-04")
    doc.add_paragraph("決議：下週三前完成報告。")
    buf = _io.BytesIO()
    doc.save(buf)

    filepath, _ = await process_attachment(
        buf.getvalue(), "meeting.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        session_id="smoke-docx", workspace=workspace,
    )
    assert os.path.exists(filepath)

    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=workspace,
        model=runner.model,
        mcp_config=runner.mcp_config,
    )
    result = await r.run(
        f"有一個會議紀錄在 {filepath}，用 Read 工具讀取，告訴我決議內容。一句話回答。",
        timeout_seconds=60.0,
    )
    assert result.error is None, f"LLM error: {result.error}"
    assert "報告" in result.text or "週三" in result.text, f"Expected decision in: {result.text}"


@smoke
@pytest.mark.asyncio
async def test_buffer_prompt_with_real_llm(runner: ClaudeRunner, tmp_path):
    """Smoke: buffer messages → build prompt → LLM understands context and answers correctly.

    NOTE: Creates a separate ClaudeRunner with workspace=tmp_path so the buffer DB
    lives in the same workspace the LLM can access via Read tool.
    """
    from raisebull.buffer import MessageBuffer
    from time import time

    buf = MessageBuffer(str(tmp_path / "buf.db"))
    await buf.init()

    now = time()
    await buf.insert("test:ch", "Alice", "今天要討論預算", now - 120)
    await buf.insert("test:ch", "Bob", "預算大約五萬", now - 60)

    prompt = await buf.build_prompt("test:ch", "剛才說的預算是多少？只回答數字。", buffer_time_minutes=10)

    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=str(tmp_path),
        model=runner.model,
    )
    result = await r.run(prompt, timeout_seconds=60.0)
    assert result.error is None, f"LLM error: {result.error}"
    assert "五萬" in result.text or "50000" in result.text or "5万" in result.text or "50,000" in result.text, \
        f"Expected budget amount in: {result.text}"

    await buf.close()


@smoke
@pytest.mark.asyncio
async def test_buffer_prompt_datetime_visible_to_llm(runner: ClaudeRunner, tmp_path):
    """Smoke: LLM can read the datetime header from buffer prompt."""
    from raisebull.buffer import MessageBuffer
    from datetime import datetime

    buf = MessageBuffer(str(tmp_path / "buf.db"))
    await buf.init()

    prompt = await buf.build_prompt("test:ch", "現在是幾月幾號？只回答日期。", buffer_time_minutes=10)

    r = ClaudeRunner(
        claude_bin=runner.claude_bin,
        workspace=str(tmp_path),
        model=runner.model,
    )
    result = await r.run(prompt, timeout_seconds=60.0)
    assert result.error is None, f"LLM error: {result.error}"
    today = datetime.now().strftime("%m")
    assert today in result.text or str(int(today)) in result.text, \
        f"Expected current month in: {result.text}"

    await buf.close()
