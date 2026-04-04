"""Integration tests for attachment processing flow.

Tests the full pipeline: bytes → router → saved file → content correct.
Uses real parsers and real filesystem, mock runner only.
"""
import os
import io
import pytest
from raisebull.parsers.router import process_attachment


class TestAttachmentFlow:
    """End-to-end attachment processing without LLM."""

    @pytest.mark.asyncio
    async def test_text_file_saved_and_readable(self, tmp_path):
        content = "Meeting notes\n\nDiscussed project timeline."
        filepath, preview = await process_attachment(
            content.encode(), "meeting-notes.txt", "text/plain",
            session_id="discord:123", workspace=str(tmp_path),
        )
        assert os.path.isfile(filepath)
        saved = open(filepath, encoding="utf-8").read()
        assert "Meeting notes" in saved
        assert "Discussed project timeline" in saved

    @pytest.mark.asyncio
    async def test_csv_rendered_as_markdown_table(self, tmp_path):
        csv_data = "product,price,qty\n高粱酒,580,10\n貢糖,120,25"
        filepath, preview = await process_attachment(
            csv_data.encode(), "inventory.csv", "text/csv",
            session_id="line:Uabc", workspace=str(tmp_path),
        )
        saved = open(filepath, encoding="utf-8").read()
        assert "|" in saved
        assert "高粱酒" in saved
        assert "580" in saved

    @pytest.mark.asyncio
    async def test_docx_paragraphs_and_tables_preserved(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("公告事項")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "項目"
        table.cell(0, 1).text = "金額"
        table.cell(1, 0).text = "材料費"
        table.cell(1, 1).text = "5000"
        buf = io.BytesIO()
        doc.save(buf)
        filepath, preview = await process_attachment(
            buf.getvalue(), "announcement.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            session_id="web:xyz", workspace=str(tmp_path),
        )
        saved = open(filepath, encoding="utf-8").read()
        assert "公告事項" in saved
        assert "材料費" in saved
        assert "5000" in saved

    @pytest.mark.asyncio
    async def test_xlsx_multi_sheet(self, tmp_path):
        import openpyxl
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Orders"
        ws1.append(["order_id", "total"])
        ws1.append(["ORD-001", 1500])
        ws2 = wb.create_sheet("Products")
        ws2.append(["name", "price"])
        ws2.append(["高粱酒", 580])
        buf = io.BytesIO()
        wb.save(buf)
        filepath, _ = await process_attachment(
            buf.getvalue(), "report.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            session_id="sess", workspace=str(tmp_path),
        )
        saved = open(filepath, encoding="utf-8").read()
        assert "Sheet: Orders" in saved
        assert "Sheet: Products" in saved
        assert "高粱酒" in saved

    @pytest.mark.asyncio
    async def test_pdf_text_extraction(self, tmp_path):
        import fitz
        # Use ASCII text long enough (>= 50 chars) to avoid the vision fallback path.
        # fitz's default font drops non-ASCII characters silently.
        page1_text = "Page 1 content - quarterly report summary with detailed notes."
        page2_text = "Page 2 content - inventory levels and shipping details here."
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), page1_text)
        page2 = doc.new_page()
        page2.insert_text((72, 72), page2_text)
        pdf_bytes = doc.tobytes()
        doc.close()
        filepath, _ = await process_attachment(
            pdf_bytes, "report.pdf", "application/pdf",
            session_id="sess", workspace=str(tmp_path),
        )
        saved = open(filepath, encoding="utf-8").read()
        assert "Page 1 content" in saved
        assert "Page 2 content" in saved

    @pytest.mark.asyncio
    async def test_image_without_vision_key(self, tmp_path):
        png = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
            b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00"
            b"\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00"
            b"\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        filepath, preview = await process_attachment(
            png, "photo.png", "image/png",
            session_id="sess", workspace=str(tmp_path),
            vision_client=None,
        )
        saved = open(filepath, encoding="utf-8").read()
        assert "vision" in saved.lower()

    @pytest.mark.asyncio
    async def test_unsupported_format_error(self, tmp_path):
        filepath, preview = await process_attachment(
            b"\x00\x01\x02", "video.mp4", "video/mp4",
            session_id="sess", workspace=str(tmp_path),
        )
        assert "不支援" in preview

    @pytest.mark.asyncio
    async def test_session_isolation(self, tmp_path):
        fp1, _ = await process_attachment(
            b"content1", "file.txt", "text/plain",
            session_id="discord:111", workspace=str(tmp_path),
        )
        fp2, _ = await process_attachment(
            b"content2", "file.txt", "text/plain",
            session_id="line:222", workspace=str(tmp_path),
        )
        assert "discord:111" in fp1 or "discord%3A111" in fp1 or os.sep + "discord:111" + os.sep in fp1
        assert "line:222" in fp2 or "line%3A222" in fp2 or os.sep + "line:222" + os.sep in fp2
        assert fp1 != fp2

    @pytest.mark.asyncio
    async def test_prompt_format(self, tmp_path):
        filepath, preview = await process_attachment(
            b"Important document content here.",
            "memo.txt", "text/plain",
            session_id="sess", workspace=str(tmp_path),
        )
        prompt = (
            f"用戶上傳了 memo.txt，已解析存放在：{filepath}\n"
            f"請用 Read 工具查看完整內容。\n"
            f"前 200 字預覽：\n{preview}"
        )
        assert filepath in prompt
        assert "Read" in prompt
        assert "Important document" in prompt
